# -*- coding: utf-8 -*-
"""
MEP Translator v2 — Dịch tài liệu kỹ thuật MEP giữ nguyên cấu trúc file gốc
(PDF lớn, Word, Excel, Ảnh/ảnh chụp màn hình) — đa ngôn ngữ, có chế độ song ngữ.

Chạy: streamlit run mep_translator.py
"""

import io
import json
import base64
import copy
import fitz  # PyMuPDF
import streamlit as st

import docx
from docx.text.paragraph import Paragraph
import openpyxl
from openpyxl.utils import range_boundaries, get_column_letter
from openpyxl.cell.cell import MergedCell

# ----------------------------------------------------------------------------
# Cấu hình chung
# ----------------------------------------------------------------------------

LANGUAGES_VI = {
    "en": "Tiếng Anh", "vi": "Tiếng Việt", "zh": "Tiếng Trung", "ja": "Tiếng Nhật",
    "ko": "Tiếng Hàn", "fr": "Tiếng Pháp", "de": "Tiếng Đức", "th": "Tiếng Thái",
    "es": "Tiếng Tây Ban Nha", "ru": "Tiếng Nga", "id": "Tiếng Indonesia",
    "ms": "Tiếng Mã Lai", "km": "Tiếng Khmer", "lo": "Tiếng Lào",
    "hi": "Tiếng Hindi", "ar": "Tiếng Ả Rập", "pt": "Tiếng Bồ Đào Nha", "it": "Tiếng Ý",
}
LANGUAGES_EN = {
    "en": "English", "vi": "Vietnamese", "zh": "Chinese", "ja": "Japanese",
    "ko": "Korean", "fr": "French", "de": "German", "th": "Thai",
    "es": "Spanish", "ru": "Russian", "id": "Indonesian", "ms": "Malay",
    "km": "Khmer", "lo": "Lao", "hi": "Hindi", "ar": "Arabic",
    "pt": "Portuguese", "it": "Italian",
}

GLOSSARY_HINT = (
    "AHU=Air Handling Unit/Bo xu ly khong khi; FCU=Fan Coil Unit; "
    "VRV/VRF=may dieu hoa trung tam; MCCB/ACB=Aptomat; "
    "tu dien=electrical panel/switchboard; ong gio=duct/air duct; "
    "ong nuoc=pipe/piping; PCCC=fire fighting/fire protection; "
    "may bom=pump; chiller=may lam lanh nuoc; thong gio=ventilation; "
    "cap thoat nuoc=water supply and drainage."
)

st.set_page_config(page_title="MEP Translator", page_icon="🛠️", layout="wide")
st.title("🛠️ MEP TRANSLATOR (Offline)")
st.caption("Dịch tài liệu kỹ thuật MEP — DOCX, PDF, Excel, Text — chạy offline bằng Hugging Face.")


# ----------------------------------------------------------------------------
# Dịch thuật (Anthropic API)
# ----------------------------------------------------------------------------

def build_system_prompt(src_label: str, tgt_label: str) -> str:
    return (
        f"Ban la chuyen gia dich thuat ky thuat chuyen nganh MEP (Co - Dien - Nuoc: "
        f"HVAC, he thong dien, cap thoat nuoc, phong chay chua chay) trong xay dung.\n"
        f"Nhiem vu: dich noi dung tu {src_label} sang {tgt_label}.\n"
        "Quy tac bat buoc:\n"
        "- Giu nguyen, KHONG dich: ma hieu thiet bi (VD: AHU-01, FCU-12, DB-3F), "
        "so hieu tieu chuan (TCVN, QCVN, ASHRAE, NFPA, ASME, SMACNA, IEC, JIS), model, "
        "ma san pham, va don vi ky thuat (kW, CFM, m3/h, Pa, mmAq, C, kVA, mm, A, V).\n"
        f"- Su dung thuat ngu MEP chuan nganh. Tham khao: {GLOSSARY_HINT}\n"
        "- Giu nguyen cau truc dinh dang, bang bieu, so thu tu va xuong dong cua "
        "van ban goc cang sat cang tot.\n"
        "- CHI xuat ra noi dung da dich. Khong them ghi chu, loi giai thich, khong "
        "lap lai van ban goc, khong dung markdown code fence."
    )
    
def translate_offline(text: str, src: str, tgt: str) -> str:
    model_name = f"Helsinki-NLP/opus-mt-{src}-{tgt}"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    inputs = tokenizer(text, return_tensors="pt", padding=True)
    translated = model.generate(**inputs)
    return tokenizer.decode(translated[0], skip_special_tokens=True)
    
# ----------------------------------------------------------------------------
# WORD (.docx) — dịch tại chỗ, giữ nguyên style/heading/bảng/bullet
# ----------------------------------------------------------------------------

def insert_paragraph_after(paragraph: Paragraph, text: str, italic=True):
    """Chèn 1 đoạn mới ngay sau paragraph, sao chép định dạng — dùng cho bản song ngữ."""
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    for run in list(new_para.runs):
        run.text = ""
    if new_para.runs:
        new_para.runs[0].text = text
        new_para.runs[0].italic = italic
        for r in new_para.runs[1:]:
            r.text = ""
    else:
        r = new_para.add_run(text)
        r.italic = italic
    return new_para


def translate_docx(client, system_prompt: str, file_bytes: bytes, bilingual: bool, progress_cb=None):
    d = docx.Document(io.BytesIO(file_bytes))

    targets = []
    for p in d.paragraphs:
        if p.text.strip():
            targets.append(p)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        targets.append(p)

    original_preview = "\n\n".join(t.text for t in targets)

    BATCH = 12
    translated_all = []
    total = max(1, (len(targets) + BATCH - 1) // BATCH)
    for i in range(0, len(targets), BATCH):
        batch = targets[i:i + BATCH]
        texts = [t.text for t in batch]
        translated_all.extend(translate_blocks_batch(client, system_prompt, texts))
        if progress_cb:
            progress_cb(min(100, int(((i // BATCH) + 1) / total * 100)))

    for para, tvalue in reversed(list(zip(targets, translated_all))):
        if bilingual:
            insert_paragraph_after(para, tvalue, italic=True)
        else:
            runs = para.runs
            if runs:
                runs[0].text = tvalue
                for r in runs[1:]:
                    r.text = ""
            else:
                para.add_run(tvalue)

    out = io.BytesIO()
    d.save(out)
    return original_preview, "\n\n".join(translated_all), out.getvalue()


# ----------------------------------------------------------------------------
# EXCEL (.xlsx) — dịch tại chỗ, giữ nguyên sheet/style/merge/formula
# ----------------------------------------------------------------------------

def insert_column_with_merge_fix(ws, insert_idx: int):
    """Chèn 1 cột tại vị trí insert_idx (1-based), tự giãn các vùng merge bị
    ảnh hưởng để không vỡ layout (ô tiêu đề merge nhiều cột, v.v.)."""
    merges = list(ws.merged_cells.ranges)
    for m in merges:
        ws.unmerge_cells(str(m))
    ws.insert_cols(insert_idx)
    for m in merges:
        min_col, min_row, max_col, max_row = range_boundaries(str(m))
        if min_col >= insert_idx:
            min_col += 1
        if max_col >= insert_idx:
            max_col += 1
        ws.merge_cells(f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{max_row}")


def copy_cell_style(src_cell, dst_cell):
    dst_cell.font = copy.copy(src_cell.font)
    dst_cell.fill = copy.copy(src_cell.fill)
    dst_cell.border = copy.copy(src_cell.border)
    dst_cell.alignment = copy.copy(src_cell.alignment)
    dst_cell.number_format = src_cell.number_format


def translate_xlsx(client, system_prompt: str, file_bytes: bytes, bilingual: bool, progress_cb=None):
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=False)
    refs = []
    preview_lines = []
    for ws in wb.worksheets:
        preview_lines.append(f"\n=== SHEET: {ws.title} ===")
        for row in ws.iter_rows():
            row_vals = []
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    refs.append({"sheet": ws.title, "coord": cell.coordinate, "text": cell.value,
                                 "row": cell.row, "col": cell.column})
                row_vals.append("" if cell.value is None else str(cell.value))
            if any(v.strip() for v in row_vals):
                preview_lines.append(" | ".join(row_vals))
    if not refs:
        raise ValueError("Khong tim thay noi dung dang chu trong bang tinh.")

    batches, cur, cur_len = [], [], 0
    for r in refs:
        l = len(r["text"]) + 4
        if cur_len + l > 1800 and cur:
            batches.append(cur)
            cur, cur_len = [], 0
        cur.append(r)
        cur_len += l
    if cur:
        batches.append(cur)

    for bi, batch in enumerate(batches):
        texts = [r["text"] for r in batch]
        translated_vals = translate_batch_strings(client, system_prompt, texts)
        for r, tval in zip(batch, translated_vals):
            r["translated"] = tval
        if progress_cb:
            progress_cb(int((bi + 1) / len(batches) * 100))

    if bilingual:
        # Song ngữ theo kiểu "cột kề cột" ngay trong sheet gốc: với mỗi cột có
        # chữ, chèn 1 cột dịch ngay bên phải, giữ nguyên cột gốc — xử lý từ
        # cột phải nhất về trái để không lệch chỉ số cột khi chèn.
        # Riêng các ô tiêu đề bị merge rộng (banner nhiều cột) không chèn cột
        # được (ô bên cạnh vẫn thuộc vùng merge) — với các ô này, gộp bản dịch
        # ngay trong cùng ô (gốc / dịch).
        for ws in wb.worksheets:
            sheet_refs = [r for r in refs if r["sheet"] == ws.title]
            if not sheet_refs:
                continue

            wide_merge_coords = set()
            for m in ws.merged_cells.ranges:
                if m.max_col > m.min_col:  # merge trải rộng nhiều cột
                    wide_merge_coords.add((m.min_row, m.min_col))

            normal_refs, wide_refs = [], []
            for r in sheet_refs:
                (wide_refs if (r["row"], r["col"]) in wide_merge_coords else normal_refs).append(r)

            # Ô tiêu đề merge rộng: gộp gốc/dịch trong cùng ô
            for r in wide_refs:
                cell = ws.cell(row=r["row"], column=r["col"])
                cell.value = f"{r['text']}  /  {r.get('translated', '')}"

            # Ô dữ liệu thường: chèn cột dịch kề bên
            cols_with_text = sorted(set(r["col"] for r in normal_refs), reverse=True)
            for col_idx in cols_with_text:
                insert_column_with_merge_fix(ws, col_idx + 1)
                src_letter = get_column_letter(col_idx)
                new_letter = get_column_letter(col_idx + 1)
                if src_letter in ws.column_dimensions:
                    ws.column_dimensions[new_letter].width = ws.column_dimensions[src_letter].width
                for r in normal_refs:
                    if r["col"] == col_idx:
                        dst_cell = ws.cell(row=r["row"], column=col_idx + 1)
                        if isinstance(dst_cell, MergedCell):
                            # an toàn dự phòng: nếu vẫn rơi vào vùng merge, gộp
                            # vào ô gốc thay vì ghi đè (tránh crash)
                            src_cell = ws.cell(row=r["row"], column=col_idx)
                            src_cell.value = f"{r['text']}  /  {r.get('translated', '')}"
                            continue
                        src_cell = ws.cell(row=r["row"], column=col_idx)
                        dst_cell.value = r.get("translated", "")
                        copy_cell_style(src_cell, dst_cell)
    else:
        for r in refs:
            wb[r["sheet"]][r["coord"]] = r.get("translated", r["text"])

    out_buf = io.BytesIO()
    wb.save(out_buf)
    original_preview = "\n".join(preview_lines).strip()
    translated_preview = "\n".join(f"{r['sheet']}!{r['coord']}: {r.get('translated','')}" for r in refs)
    return original_preview, translated_preview, out_buf.getvalue()


# ----------------------------------------------------------------------------
# PDF (kể cả PDF lớn) — trích xuất theo block/trang, xuất ra DOCX giữ đúng
# thứ tự đọc + ngắt trang gốc (không ghi đè trực tiếp lên PDF để tránh lỗi
# hiển thị dấu tiếng Việt / chữ Hán / chữ Nhật do thiếu font nhúng sẵn).
# ----------------------------------------------------------------------------

def get_pdf_page_count(file_bytes: bytes) -> int:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    n = doc.page_count
    doc.close()
    return n


def extract_pdf_blocks(file_bytes: bytes, page_from: int, page_to: int):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_blocks = []
    for pno in range(page_from - 1, page_to):
        page = doc.load_page(pno)
        blocks = page.get_text("blocks")
        blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
        texts = [b[4].strip() for b in blocks if b[4] and b[4].strip()]
        if texts:
            pages_blocks.append((pno + 1, texts))
    doc.close()
    if not pages_blocks:
        raise ValueError("Khong trich xuat duoc van ban trong khoang trang da chon "
                          "(co the la PDF dang scan anh — hay chup man hinh trang do "
                          "va dung che do Anh thay the).")
    return pages_blocks


def translate_pdf_to_docx(client, system_prompt: str, file_bytes: bytes,
                           page_from: int, page_to: int, bilingual: bool, progress_cb=None):
    pages_blocks = extract_pdf_blocks(file_bytes, page_from, page_to)

    out_doc = docx.Document()
    out_doc.add_heading("Bản dịch tài liệu PDF (MEP Translator)", level=1)

    original_parts, translated_parts = [], []
    total_pages = len(pages_blocks)
    for idx, (pno, texts) in enumerate(pages_blocks):
        out_doc.add_heading(f"Trang {pno}", level=2)
        translated_texts = translate_blocks_batch(client, system_prompt, texts)
        for orig, tval in zip(texts, translated_texts):
            if bilingual:
                p1 = out_doc.add_paragraph(orig)
                if p1.runs:
                    p1.runs[0].italic = True
                out_doc.add_paragraph(tval)
            else:
                out_doc.add_paragraph(tval)
        original_parts.append(f"--- Trang {pno} ---\n" + "\n\n".join(texts))
        translated_parts.append(f"--- Trang {pno} ---\n" + "\n\n".join(translated_texts))
        if progress_cb:
            progress_cb(int((idx + 1) / total_pages * 100))
        if idx < total_pages - 1:
            out_doc.add_page_break()

    out = io.BytesIO()
    out_doc.save(out)
    return "\n\n".join(original_parts), "\n\n".join(translated_parts), out.getvalue()


# ----------------------------------------------------------------------------
# Giao diện Streamlit
# ----------------------------------------------------------------------------

with st.sidebar:
    st.header("Cấu hình dịch")
    src_lang = st.selectbox("Ngôn ngữ nguồn", list(LANGUAGES_VI.keys()), format_func=lambda x: LANGUAGES_VI[x])
    tgt_lang = st.selectbox("Ngôn ngữ đích", list(LANGUAGES_VI.keys()), format_func=lambda x: LANGUAGES_VI[x])

uploaded = st.file_uploader("Tải lên tệp cần dịch (PDF, DOCX, XLSX, TXT)", type=["pdf","docx","xlsx","txt"])

if uploaded:
    kind = uploaded.name.lower().split(".")[-1]
    file_bytes = uploaded.getvalue()

    if st.button("🌐 Dịch tài liệu"):
        if kind == "docx":
            out_bytes = translate_docx(file_bytes, src_lang, tgt_lang)
            st.download_button("⬇️ Tải DOCX đã dịch", out_bytes, file_name="translated.docx")
        elif kind == "xlsx":
            out_bytes = translate_xlsx(file_bytes, src_lang, tgt_lang)
            st.download_button("⬇️ Tải Excel đã dịch", out_bytes, file_name="translated.xlsx")
        elif kind == "pdf":
            out_bytes = translate_pdf(file_bytes, src_lang, tgt_lang)
            st.download_button("⬇️ Tải DOCX từ PDF đã dịch", out_bytes, file_name="translated_from_pdf.docx")
        elif kind == "txt":
            text = file_bytes.decode("utf-8")
            translated = translate_offline(text, src_lang, tgt_lang)
            st.text_area("Bản dịch", translated, height=400)
            st.download_button("⬇️ Tải TXT đã dịch", translated, file_name="translated.txt")

